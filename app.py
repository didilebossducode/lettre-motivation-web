from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import tempfile
from pathlib import Path
import json

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', os.path.join(tempfile.gettempdir(), 'lettre_motivation_uploads'))
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Assurez-vous que le dossier uploads existe
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def cleanup_old_files():
    """Nettoie les fichiers de plus de 1 heure"""
    now = datetime.now()
    for file in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
        if os.path.isfile(file_path):
            file_time = datetime.fromtimestamp(os.path.getctime(file_path))
            if (now - file_time).total_seconds() > 3600:  # 1 heure
                os.remove(file_path)

def create_letter(company_name, job_title, duration, start_date, today_date, custom_text):
    cleanup_old_files()
    
    doc = Document()
    
    # Configuration de la page
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)  # A4
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(1.18)
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Ajouter la date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.add_run(today_date)
    
    # Ajouter deux lignes vides
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Objet de la lettre
    object_paragraph = doc.add_paragraph()
    object_paragraph.add_run('Objet : ').bold = True
    object_paragraph.add_run(f'Candidature au poste de {job_title} chez {company_name}')
    
    # Ajouter une ligne vide
    doc.add_paragraph()
    
    # Formule de politesse
    doc.add_paragraph('Madame, Monsieur,')
    
    # Corps de la lettre
    body = doc.add_paragraph()
    body.add_run(f'Je me permets de vous adresser ma candidature pour le poste de {job_title} au sein de votre entreprise {company_name}')
    if duration:
        body.add_run(f' pour une durée de {duration}')
    if start_date:
        body.add_run(f', à partir du {start_date}')
    body.add_run('.')
    
    # Paragraphe personnalisé
    if custom_text:
        doc.add_paragraph()
        doc.add_paragraph(custom_text)
    
    # Formule de fin
    doc.add_paragraph()
    doc.add_paragraph('Je me tiens à votre disposition pour un entretien et vous prie d\'agréer, Madame, Monsieur, l\'expression de mes salutations distinguées.')
    
    # Créer un nom de fichier unique
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f"lettre_motivation_{timestamp}.docx"
    
    # Chemin complet
    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
    
    # Sauvegarder le document
    doc.save(docx_path)
    return docx_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_letter():
    try:
        data = request.get_json()
        
        # Générer la lettre
        output_path = create_letter(
            company_name=data.get('company_name', ''),
            job_title=data.get('job_title', ''),
            duration=data.get('duration', ''),
            start_date=data.get('start_date', ''),
            today_date=data.get('today_date', ''),
            custom_text=data.get('custom_text', '')
        )
        
        # Retourner le chemin du fichier
        return jsonify({
            'success': True,
            'file_path': output_path
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_file(
            os.path.join(app.config['UPLOAD_FOLDER'], filename),
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return str(e), 404

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
